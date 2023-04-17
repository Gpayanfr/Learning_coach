import docx
import os 


def CompleterDecimaux(liste,plan_perso) : 
   os.chdir('Banque_exercices\\Nombres_et_calcul\\Decimaux\\PreRequis')
   if liste[1] == '1' :
      plan_perso.add_picture('Decimaux1_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[2] == '1' :
      plan_perso.add_picture('Decimaux2_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[3] == '1' :
      plan_perso.add_picture('Decimaux3_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[4] == '1' :
      plan_perso.add_picture('Decimaux4_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[5] == '1' :
      plan_perso.add_picture('Decimaux5_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[6] == '1' :
      plan_perso.add_picture('Decimaux6_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[7] == '1' :
      plan_perso.add_picture('Decimaux7_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[8] == '1' :
      plan_perso.add_picture('Decimaux8_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[9] == '1' :
      plan_perso.add_picture('Decimaux9_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')

   os.chdir('Remediation')
   if liste[1] == '2' :
      plan_perso.add_picture('Decimaux1_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[2] == '2' :
      plan_perso.add_picture('Decimaux2_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[3] == '2' :
      plan_perso.add_picture('Decimaux3_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[4] == '2' :
      plan_perso.add_picture('Decimaux4_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[5] == '2' :
      plan_perso.add_picture('Decimaux5_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[6] == '2' :
      plan_perso.add_picture('Decimaux6_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[7] == '2' :
      plan_perso.add_picture('Decimaux7_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[8] == '2' :
      plan_perso.add_picture('Decimaux8_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[9] == '2' :
      plan_perso.add_picture('Decimaux9_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')

   os.chdir('Consolidation')
   if liste[1] == '3' :
      plan_perso.add_picture('Decimaux1_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[2] == '3' :
      plan_perso.add_picture('Decimaux2_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[3] == '3' :
      plan_perso.add_picture('Decimaux3_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[4] == '3' :
      plan_perso.add_picture('Decimaux4_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[5] == '3' :
      plan_perso.add_picture('Decimaux5_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[6] == '3' :
      plan_perso.add_picture('Decimaux6_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[7] == '3' :
      plan_perso.add_picture('Decimaux7_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[8] == '3' :
      plan_perso.add_picture('Decimaux8_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[9] == '3' :
      plan_perso.add_picture('Decimaux9_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')
   
   os.chdir('Approfondissement')
   if liste[1] == '4' :
      plan_perso.add_picture('Decimaux1_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[2] == '4' :
      plan_perso.add_picture('Decimaux2_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[3] == '4' :
      plan_perso.add_picture('Decimaux3_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[4] == '4' :
      plan_perso.add_picture('Decimaux4_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[5] == '4' :
      plan_perso.add_picture('Decimaux5_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[6] == '4' :
      plan_perso.add_picture('Decimaux6_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[7] == '4' :
      plan_perso.add_picture('Decimaux7_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[8] == '4' :
      plan_perso.add_picture('Decimaux8_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[9] == '4' :
      plan_perso.add_picture('Decimaux9_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
      
   os.chdir('..')
   os.chdir('..')
   os.chdir('..')
   os.chdir('..') 
