import docx
import os 

def CompleterFractions(liste,plan_perso) : 
    os.chdir('Banque_exercices\\Nombres_et_calcul\\Fractions\\PreRequis') 
    if liste[13] == '1' :
        plan_perso.add_picture('Fractions1_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[14] == '1' :
        plan_perso.add_picture('Fractions2_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[15] == '1' :
        plan_perso.add_picture('Fractions3_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[16] == '1' :
        plan_perso.add_picture('Fractions4_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[17] == '1' :
        plan_perso.add_picture('Fractions5_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[18] == '1' :
        plan_perso.add_picture('Fractions6_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[19] == '1' :
        plan_perso.add_picture('Fractions7_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[20] == '1' :
        plan_perso.add_picture('Fractions8_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[21] == '1' :
        plan_perso.add_picture('Fractions9_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    os.chdir('..')

    os.chdir('Remediation')
    if liste[13] == '2' :
        plan_perso.add_picture('Fractions1_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[14] == '2' :
        plan_perso.add_picture('Fractions2_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[15] == '2' :
        plan_perso.add_picture('Fractions3_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[16] == '2' :
        plan_perso.add_picture('Fractions4_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[17] == '2' :
        plan_perso.add_picture('Fractions5_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[18] == '2' :
        plan_perso.add_picture('Fractions6_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[19] == '2' :
        plan_perso.add_picture('Fractions7_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[20] == '2' :
        plan_perso.add_picture('Fractions8_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[21] == '2' :
        plan_perso.add_picture('Fractions9_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    os.chdir('..')

    os.chdir('Consolidation')
    if liste[13] == '3' :
        plan_perso.add_picture('Fractions1_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[14] == '3' :
        plan_perso.add_picture('Fractions2_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[15] == '3' :
        plan_perso.add_picture('Fractions3_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[16] == '3' :
        plan_perso.add_picture('Fractions4_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[17] == '3' :
        plan_perso.add_picture('Fractions5_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[18] == '3' :
        plan_perso.add_picture('Fractions6_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[19] == '3' :
        plan_perso.add_picture('Fractions7_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[20] == '3' :
        plan_perso.add_picture('Fractions8_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[21] == '3' :
        plan_perso.add_picture('Fractions9_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    os.chdir('..')

    os.chdir('Approfondissement')
    if liste[13] == '4' :
        plan_perso.add_picture('Fractions1_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[14] == '4' :
        plan_perso.add_picture('Fractions2_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[15] == '4' :
        plan_perso.add_picture('Fractions3_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[16] == '4' :
        plan_perso.add_picture('Fractions4_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[17] == '4' :
        plan_perso.add_picture('Fractions5_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[18] == '4' :
        plan_perso.add_picture('Fractions6_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[19] == '4' :
        plan_perso.add_picture('Fractions7_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[20] == '4' :
        plan_perso.add_picture('Fractions8_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
    if liste[21] == '4' :
        plan_perso.add_picture('Fractions9_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
        
    os.chdir('..')
    os.chdir('..')
    os.chdir('..')
    os.chdir('..') 