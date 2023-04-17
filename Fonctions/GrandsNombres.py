import docx
import os 


def CompleterGrandsNombres(liste,plan_perso) : 
   os.chdir('Banque_exercices\\Nombres_et_calcul\\GrandsNombres\\PreRequis') 
   if liste[10] == '1' :
      plan_perso.add_picture('GrandsNombres1_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[11] == '1' :
      plan_perso.add_picture('GrandsNombres2_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[12] == '1' :
      plan_perso.add_picture('GrandsNombres3_PreRequis.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')

   os.chdir('Remediation')
   if liste[10] == '2' :
      plan_perso.add_picture('GrandsNombres1_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[11] == '2' :
      plan_perso.add_picture('GrandsNombres2_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[12] == '2' :
      plan_perso.add_picture('GrandsNombres3_Remediation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')

   os.chdir('Consolidation')
   if liste[10] == '3' :
      plan_perso.add_picture('GrandsNombres1_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[11] == '3' :
      plan_perso.add_picture('GrandsNombres2_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[12] == '3' :
      plan_perso.add_picture('GrandsNombres3_Consolidation.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   os.chdir('..')

   os.chdir('Approfondissement')
   if liste[10] == '4' :
      plan_perso.add_picture('GrandsNombres1_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[11] == '4' :
      plan_perso.add_picture('GrandsNombres2_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))
   if liste[12] == '4' :
      plan_perso.add_picture('GrandsNombres3_Approfondissement.png', width=docx.shared.Cm(18), height=docx.shared.Cm(12))

   os.chdir('..')
   os.chdir('..')
   os.chdir('..')
   os.chdir('..') 